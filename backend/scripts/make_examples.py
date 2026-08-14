"""生成示例文档（markdown/txt/docx/xlsx），用于快速演示。

用法: python scripts/make_examples.py
"""
from __future__ import annotations

from pathlib import Path

EXAMPLES = Path(__file__).resolve().parents[2] / "examples"


def main() -> None:
    EXAMPLES.mkdir(parents=True, exist_ok=True)

    # 1) Markdown 员工手册
    (EXAMPLES / "员工手册.md").write_text(
        """# 星辰科技员工手册（2025 版）

## 第一章 总则

本手册适用于星辰科技全体正式员工，是公司日常管理的基本依据。

## 第二章 考勤制度

- 工作时间：周一至周五 9:00–18:00，午休 12:00–13:00。
- 员工须通过企业微信打卡，迟到 30 分钟以上按事假处理。
- 每月迟到累计超过 3 次，将影响当月绩效评级。

## 第三章 休假制度

- 年假：入职满 1 年享受 5 天，满 3 年 7 天，满 5 年 10 天。
- 病假需提供三甲医院证明；婚假 3 天；产假按国家规定执行。
- 事假须提前 1 天在 OA 系统申请，由直属主管审批。

## 第四章 信息安全

- 员工须定期修改密码，禁止将账号借与他人使用。
- 涉及客户数据的文件不得通过个人网盘传输。
- 离职时须归还全部设备与资料，并完成权限回收流程。

## 第五章 绩效考核

- 绩效按季度评估，由主管评分 + 同事互评构成。
- 连续两个季度考核为 C 的员工将进入辅导计划。
- 年度优秀员工将获得额外奖金与晋升优先权。
""",
        encoding="utf-8",
    )

    # 2) TXT 产品 FAQ
    (EXAMPLES / "产品FAQ.txt").write_text(
        """星辰云平台产品常见问题（FAQ）

Q1: 星辰云平台如何计费？
A: 平台按套餐订阅制收费，分为基础版（999元/月）、专业版（2999元/月）、旗舰版（8999元/月）。旗舰版支持私有化部署。

Q2: 数据存储在哪里？
A: 公有云版本数据存储于国内合规机房，支持数据加密；旗舰版可部署在客户自有机房或公有云 VPC。

Q3: 是否支持单点登录（SSO）？
A: 专业版及以上支持 SAML 2.0 与 OIDC 协议的企业 SSO 对接。

Q4: 如何进行数据迁移？
A: 提供标准 CSV/API 双通道迁移工具，迁移期间业务不中断，全程有技术支持护航。

Q5: 售后服务如何保障？
A: 专业版提供 5×8 小时服务，旗舰版提供 7×24 小时专属服务群，响应时效 15 分钟。

Q6: 免费试用？
A: 所有版本均提供 14 天免费试用，无需绑定银行卡。
""",
        encoding="utf-8",
    )

    # 3) docx 制度文件
    from docx import Document

    doc = Document()
    doc.add_heading("差旅报销管理制度", level=1)
    doc.add_paragraph("一、适用范围：全体员工因公出差发生的交通、住宿、餐饮费用。")
    doc.add_paragraph("二、报销标准：一线城市住宿上限 500 元/晚，二线城市 350 元/晚；高铁二等座实报实销。")
    doc.add_paragraph("三、审批流程：出差前在 OA 提交申请，报销单由部门经理审批后 5 个工作日内打款。")
    doc.add_paragraph("四、发票要求：必须提供增值税普通发票或专用发票，无发票不予报销。")
    doc.save(str(EXAMPLES / "差旅报销管理制度.docx"))

    # 4) xlsx 数据
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "客户信息"
    ws.append(["客户名称", "行业", "套餐", "签约金额(万)", "续约率"])
    ws.append(["华信集团", "制造业", "旗舰版", 120, "95%"])
    ws.append(["云途物流", "物流", "专业版", 60, "88%"])
    ws.append(["金科股份", "金融", "旗舰版", 200, "97%"])
    ws.append(["绿洲教育", "教育", "基础版", 20, "76%"])
    wb.save(str(EXAMPLES / "客户数据.xlsx"))

    print(f"示例文档已生成到: {EXAMPLES}")
    for p in sorted(EXAMPLES.iterdir()):
        print("  -", p.name)


if __name__ == "__main__":
    main()
